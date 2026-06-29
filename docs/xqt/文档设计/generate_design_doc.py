#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
基于 docs/xqt/文档设计/设计文档大纲_v3.md 生成 .docx 模板。
- 仅输出结构与格式占位，不填充具体内容
- 重点：标题层级、页面布局、字体行距、图/表占位、篇幅分配标注
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

# ============= 1. 文档级配置 =============
OUTPUT_PATH = "D:/PClearning/AgentProjects/softbei/docs/xqt/文档设计/项目设计文档_模板_v3.2.docx"

# A4 页面 + 默认中文论文页边距
PAGE_W_MM = 210
PAGE_H_MM = 297
MARGIN_TOP_CM = 2.54
MARGIN_BOTTOM_CM = 2.54
MARGIN_LEFT_CM = 3.18
MARGIN_RIGHT_CM = 3.18

# 中文字体（按 CLAUDE.md：正文宋体小四 + 1.5 倍行距）
BODY_FONT_ZH = "宋体"
BODY_FONT_EN = "Times New Roman"
HEADING_FONT_ZH = "黑体"
HEADING_FONT_EN = "Times New Roman"
CODE_FONT = "Consolas"

# 字号（pt）：小四=12pt、五号=10.5pt、小三=15pt、四号=14pt、小二=18pt、二号=22pt、小一=24pt
SZ_TITLE = 22       # 一级标题（封面主标题）
SZ_PART = 18        # 部分标题（"第一部分 …"）
SZ_H1 = 16          # 章标题（"1.1 …"）
SZ_H2 = 14          # 节标题
SZ_H3 = 13          # 小节标题
SZ_BODY = 12        # 正文（小四）
SZ_SMALL = 10.5     # 辅助说明（五号）
SZ_TABLE = 10.5     # 表格内文字

LINE_SPACING_BODY = 1.5
LINE_SPACING_TABLE = 1.0

# 颜色（用于占位提示）
GRAY_LIGHT = RGBColor(0x99, 0x99, 0x99)   # 占位符浅灰
GRAY_MID = RGBColor(0x66, 0x66, 0x66)     # 元信息中灰
ACCENT = RGBColor(0x1F, 0x4E, 0x79)       # 强调深蓝（页眉/页脚）
TABLE_HEADER = RGBColor(0xD9, 0xE2, 0xF3)  # 表头浅蓝

# ============= 2. 文档初始化 =============
doc = Document()


def set_cn_font(run, font_zh=BODY_FONT_ZH, font_en=BODY_FONT_EN, size=SZ_BODY,
                bold=False, italic=False, color=None):
    """设置 run 的中英文字体与字号。"""
    run.font.name = font_en
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_zh)
    rFonts.set(qn("w:ascii"), font_en)
    rFonts.set(qn("w:hAnsi"), font_en)
    rFonts.set(qn("w:cs"), font_en)


def set_paragraph_format(paragraph, line_spacing=LINE_SPACING_BODY,
                         space_before=0, space_after=0,
                         alignment=WD_ALIGN_PARAGRAPH.LEFT,
                         first_line_indent=None,
                         keep_with_next=False):
    pf = paragraph.paragraph_format
    pf.line_spacing = line_spacing
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.alignment = alignment
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    pf.keep_with_next = keep_with_next


def add_paragraph(text="", size=SZ_BODY, bold=False, italic=False,
                  color=None, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                  first_line_indent=None, line_spacing=LINE_SPACING_BODY,
                  space_before=0, space_after=0, keep_with_next=False,
                  font_zh=BODY_FONT_ZH, font_en=BODY_FONT_EN,
                  style=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    set_paragraph_format(p, line_spacing, space_before, space_after,
                         alignment, first_line_indent, keep_with_next)
    if text:
        run = p.add_run(text)
        set_cn_font(run, font_zh, font_en, size, bold, italic, color)
    return p


def add_heading(text, level, page_break_before=False):
    """自定义标题：page_break_before 控制是否分页。"""
    if page_break_before:
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_break(WD_BREAK.PAGE)
        p = doc.add_paragraph()
    else:
        p = doc.add_paragraph()

    # 字体与字号按层级
    if level == 0:  # 部分（"第一部分 …"）
        size, bold, font_zh = SZ_PART, True, HEADING_FONT_ZH
        space_before, space_after, align = 24, 18, WD_ALIGN_PARAGRAPH.CENTER
    elif level == 1:  # 一级标题（"1.1 …"）
        size, bold, font_zh = SZ_H1, True, HEADING_FONT_ZH
        space_before, space_after, align = 18, 12, WD_ALIGN_PARAGRAPH.LEFT
    elif level == 2:  # 二级标题
        size, bold, font_zh = SZ_H2, True, HEADING_FONT_ZH
        space_before, space_after, align = 12, 6, WD_ALIGN_PARAGRAPH.LEFT
    else:  # 三级标题
        size, bold, font_zh = SZ_H3, True, HEADING_FONT_ZH
        space_before, space_after, align = 6, 3, WD_ALIGN_PARAGRAPH.LEFT

    set_paragraph_format(p, LINE_SPACING_BODY, space_before, space_after,
                         align, None, True)
    run = p.add_run(text)
    set_cn_font(run, font_zh, HEADING_FONT_EN, size, bold)
    # 标记为大纲级别，TOC 才会收录
    pPr = p._element.get_or_add_pPr()
    outline = OxmlElement("w:outlineLvl")
    outline.set(qn("w:val"), str(level))
    pPr.append(outline)
    return p


def add_placeholder(label, hint="", height_lines=4):
    """插入「[待填充]」灰色占位段 + 高度占位空行。"""
    # 占位标签行
    p = doc.add_paragraph()
    set_paragraph_format(p, LINE_SPACING_BODY, 6, 3, WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run("[待填充] " + label)
    set_cn_font(run, BODY_FONT_ZH, BODY_FONT_EN, SZ_BODY,
                italic=True, color=GRAY_LIGHT)
    if hint:
        run2 = p.add_run("  · " + hint)
        set_cn_font(run2, BODY_FONT_ZH, BODY_FONT_EN, SZ_SMALL,
                    italic=True, color=GRAY_MID)
    # 高度占位（若干空行，让后续填写时有视觉空间）
    for _ in range(height_lines):
        ep = doc.add_paragraph()
        set_paragraph_format(ep, LINE_SPACING_BODY, 0, 0,
                             WD_ALIGN_PARAGRAPH.LEFT)


def add_figure_placeholder(figure_id, caption, height_cm=6.0):
    """图占位框（带图编号和图标题）。"""
    p = doc.add_paragraph()
    set_paragraph_format(p, LINE_SPACING_BODY, 12, 6,
                         WD_ALIGN_PARAGRAPH.CENTER, keep_with_next=True)
    run = p.add_run("【图占位】 ")
    set_cn_font(run, BODY_FONT_ZH, BODY_FONT_EN, SZ_BODY,
                bold=True, color=GRAY_LIGHT)
    run2 = p.add_run(figure_id + "  " + caption)
    set_cn_font(run2, BODY_FONT_ZH, BODY_FONT_EN, SZ_BODY,
                bold=True, color=GRAY_MID)
    # 灰色虚线占位框（用空段 + 下边框模拟）
    box = doc.add_paragraph()
    set_paragraph_format(box, 1.0, 0, 6, WD_ALIGN_PARAGRAPH.CENTER,
                         keep_with_next=True)
    pPr = box._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for edge in ["top", "left", "bottom", "right"]:
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "dashed")
        b.set(qn("w:sz"), "6")
        b.set(qn("w:space"), "4")
        b.set(qn("w:color"), "BFBFBF")
        pBdr.append(b)
    pPr.append(pBdr)
    # 占位框内文字
    run = box.add_run("（此处插入 " + figure_id + " · 约 " + str(int(height_cm)) + " cm 高）")
    set_cn_font(run, BODY_FONT_ZH, BODY_FONT_EN, SZ_SMALL,
                italic=True, color=GRAY_LIGHT)
    # 设置占位框最小高度
    box.paragraph_format.space_before = Pt(height_cm * 8)
    # 图下说明（提问句）
    note = doc.add_paragraph()
    set_paragraph_format(note, LINE_SPACING_BODY, 0, 12,
                         WD_ALIGN_PARAGRAPH.LEFT)
    run = note.add_run("图说（本图要回答的问题）：__________________________")
    set_cn_font(run, BODY_FONT_ZH, BODY_FONT_EN, SZ_SMALL,
                italic=True, color=GRAY_MID)


def add_table_placeholder(table_id, header, rows_hint, caption):
    """表占位（表前 1 段 + 表 + 表后 1 段结论）。"""
    # 表前说明
    pre = doc.add_paragraph()
    set_paragraph_format(pre, LINE_SPACING_BODY, 6, 3,
                         WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Cm(0.74))
    run = pre.add_run("（表前说明 · 1 段）本表要回答的问题 / 选取维度：__________________________")
    set_cn_font(run, BODY_FONT_ZH, BODY_FONT_EN, SZ_BODY,
                italic=True, color=GRAY_MID)
    # 表标题
    p = doc.add_paragraph()
    set_paragraph_format(p, LINE_SPACING_BODY, 6, 3,
                         WD_ALIGN_PARAGRAPH.CENTER, keep_with_next=True)
    run = p.add_run("表 " + table_id + "  " + caption)
    set_cn_font(run, HEADING_FONT_ZH, HEADING_FONT_EN, SZ_BODY, bold=True)
    # 表内容
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Light Grid Accent 1"
    table.autofit = True
    hdr = table.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        set_paragraph_format(p, LINE_SPACING_TABLE, 0, 0,
                             WD_ALIGN_PARAGRAPH.CENTER)
        r = p.add_run(h)
        set_cn_font(r, HEADING_FONT_ZH, HEADING_FONT_EN, SZ_TABLE, bold=True)
        # 表头底色
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "D9E2F3")
        tcPr.append(shd)
    # 占位行
    for _ in range(rows_hint):
        row = table.add_row()
        for cell in row.cells:
            cell.text = ""
            p = cell.paragraphs[0]
            set_paragraph_format(p, LINE_SPACING_TABLE, 0, 0,
                                 WD_ALIGN_PARAGRAPH.LEFT)
            r = p.add_run("________")
            set_cn_font(r, BODY_FONT_ZH, BODY_FONT_EN, SZ_TABLE,
                        color=GRAY_LIGHT)
    # 表后结论
    post = doc.add_paragraph()
    set_paragraph_format(post, LINE_SPACING_BODY, 6, 12,
                         WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Cm(0.74))
    run = post.add_run("（表后结论 · 1-2 句）由表可见：__________________________")
    set_cn_font(run, BODY_FONT_ZH, BODY_FONT_EN, SZ_BODY,
                italic=True, color=GRAY_MID)


def add_target_box(pages_target, priority="中", note=""):
    """篇幅目标标注框（每节开头或末尾插入）。"""
    p = doc.add_paragraph()
    set_paragraph_format(p, 1.0, 3, 3, WD_ALIGN_PARAGRAPH.RIGHT)
    run = p.add_run(f"[篇幅目标: {pages_target} 页]  [优先级: {priority}]")
    set_cn_font(run, BODY_FONT_ZH, BODY_FONT_EN, SZ_SMALL,
                italic=True, color=GRAY_MID)
    if note:
        run2 = p.add_run(f"  · {note}")
        set_cn_font(run2, BODY_FONT_ZH, BODY_FONT_EN, SZ_SMALL,
                    italic=True, color=GRAY_MID)


def add_meta_line(text):
    """元信息行（灰色、小号）。"""
    p = doc.add_paragraph()
    set_paragraph_format(p, 1.0, 0, 0, WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run(text)
    set_cn_font(run, BODY_FONT_ZH, BODY_FONT_EN, SZ_SMALL,
                italic=True, color=GRAY_MID)


def add_page_break():
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


# ============= 3. 页面设置 =============
section = doc.sections[0]
section.page_height = Mm(PAGE_H_MM)
section.page_width = Mm(PAGE_W_MM)
section.top_margin = Cm(MARGIN_TOP_CM)
section.bottom_margin = Cm(MARGIN_BOTTOM_CM)
section.left_margin = Cm(MARGIN_LEFT_CM)
section.right_margin = Cm(MARGIN_RIGHT_CM)
section.header_distance = Cm(1.5)
section.footer_distance = Cm(1.75)

# 页眉（区分奇偶：奇数页右对齐，偶数页左对齐）
header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hr = hp.add_run("智学工坊 · 个性化资源生成与学习多智能体系统 · 项目设计文档 v3.2")
set_cn_font(hr, HEADING_FONT_ZH, HEADING_FONT_EN, SZ_SMALL,
            color=ACCENT)

# 页脚（页码：奇数页右、偶数页左）
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run("— ")
set_cn_font(fr, BODY_FONT_ZH, BODY_FONT_EN, SZ_SMALL, color=ACCENT)
# 域代码：插入页码
fld_run = fp.add_run()
fldChar1 = OxmlElement("w:fldChar")
fldChar1.set(qn("w:fldCharType"), "begin")
instrText = OxmlElement("w:instrText")
instrText.text = " PAGE "
fldChar2 = OxmlElement("w:fldChar")
fldChar2.set(qn("w:fldCharType"), "end")
fld_run._r.append(fldChar1)
fld_run._r.append(instrText)
fld_run._r.append(fldChar2)
set_cn_font(fld_run, BODY_FONT_ZH, BODY_FONT_EN, SZ_SMALL, color=ACCENT)
fr2 = fp.add_run(" —")
set_cn_font(fr2, BODY_FONT_ZH, BODY_FONT_EN, SZ_SMALL, color=ACCENT)


# ============= 4. 封面 =============
def build_cover():
    # 上方留白
    for _ in range(4):
        add_paragraph("", size=SZ_BODY)

    # 项目名（主标题）
    add_paragraph("", size=SZ_BODY)
    p = doc.add_paragraph()
    set_paragraph_format(p, 1.5, 12, 12, WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("个性化资源生成与学习多智能体系统")
    set_cn_font(r, HEADING_FONT_ZH, HEADING_FONT_EN, SZ_TITLE, bold=True)

    # 英文名
    p = doc.add_paragraph()
    set_paragraph_format(p, 1.5, 0, 6, WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("AI Study Companion: A Multi-Agent System for Personalized Learning Resource Generation")
    set_cn_font(r, BODY_FONT_ZH, HEADING_FONT_EN, SZ_H1, italic=True)

    # 副标题（产品品牌）
    p = doc.add_paragraph()
    set_paragraph_format(p, 1.5, 18, 6, WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("— 智学工坊 —")
    set_cn_font(r, HEADING_FONT_ZH, HEADING_FONT_EN, SZ_H2, bold=True)

    # 留白
    for _ in range(3):
        add_paragraph("", size=SZ_BODY)

    # 赛题信息
    info_lines = [
        ("赛题", "第十五届中国软件杯 A3 赛题（科大讯飞股份有限公司）"),
        ("项目", "基于大模型的个性化资源生成与学习多智能体系统"),
        ("版本", "项目设计文档 v3.2 · 评审批注闭环版"),
        ("日期", "2026 年 06 月 29 日"),
        ("学校", "电子科技大学（UESTC）"),
        ("作者", "答辩团队（按实际填写）"),
        ("指导教师", "按实际填写"),
    ]
    for label, value in info_lines:
        p = doc.add_paragraph()
        set_paragraph_format(p, 1.5, 0, 0, WD_ALIGN_PARAGRAPH.CENTER)
        r1 = p.add_run(label + "：")
        set_cn_font(r1, HEADING_FONT_ZH, HEADING_FONT_EN, SZ_H3, bold=True)
        r2 = p.add_run(value)
        set_cn_font(r2, BODY_FONT_ZH, BODY_FONT_EN, SZ_BODY)

    # 底部留白
    for _ in range(4):
        add_paragraph("", size=SZ_BODY)

    p = doc.add_paragraph()
    set_paragraph_format(p, 1.0, 12, 0, WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("配套答辩 PPT：智学工坊 · UESTC · A3 Thesis Defense（17 张）")
    set_cn_font(r, BODY_FONT_ZH, BODY_FONT_EN, SZ_SMALL,
                italic=True, color=GRAY_MID)
    p = doc.add_paragraph()
    set_paragraph_format(p, 1.0, 0, 0, WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("本模板为 v3.2 结构骨架，章节内容待填充")
    set_cn_font(r, BODY_FONT_ZH, BODY_FONT_EN, SZ_SMALL,
                italic=True, color=GRAY_LIGHT)

    add_page_break()


# ============= 5. 文档编写说明（前言） =============
def build_front_matter():
    add_heading("文档编写说明", level=0)

    add_meta_line("本文档为「智学工坊」项目设计文档的 v3.2 模板骨架。")
    add_meta_line("正文部分按 6 大版块 / 67-71 页布局，附录 8-12 页。")
    add_meta_line("所有 [待填充] 标记处均为待撰写内容占位。")
    add_paragraph("", size=SZ_SMALL)

    # 1. 叙事主线
    add_heading("1. 叙事主线", level=1)
    add_placeholder("（约 1 段）", hint="六板块自上而下展开的逻辑线", height_lines=4)

    # 2. 四层架构说明
    add_heading("2. 四层架构", level=1)
    add_placeholder("（约 1 段）", hint="展示层 / 智能体路由总线 / 智能体功能层 / 基础能力层", height_lines=3)

    # 3. 创新点结构
    add_heading("3. 创新点结构（v3.2 新增）", level=1)
    add_placeholder("（约 1 段）", hint="3 PPT 创新点 + 1 文档独有", height_lines=3)

    # 4. 12 智能体统一口径
    add_heading("4. 12 智能体统一口径", level=1)
    add_placeholder("（约 1 段）", hint="profile / planner / clarify / doc / mindmap / quiz / code / summary / anim / kg / safety / recommend", height_lines=3)

    # 5. 五层防幻觉
    add_heading("5. 五层递进防幻觉架构", level=1)
    add_placeholder("（约 1 段）", hint="第一到第五层（与答辩 PPT 一致）", height_lines=3)

    # 6. 实现状态
    add_heading("6. 实现状态", level=1)
    add_placeholder("（约 1 段）", hint="所有实现状态以代码仓库当前状态为准", height_lines=2)

    # 7. AI 工具合规
    add_heading("7. AI 工具合规", level=1)
    add_placeholder("（约 1 段）", hint="附录 D 披露讯飞相关 AI 辅助工具", height_lines=2)

    add_page_break()

    # 写作风格与排版规范（v3.2 新增）
    add_heading("写作风格与排版规范（v3.2 新增）", level=1)
    add_meta_line("落实 YC 2026.06.27 评审批注 #0 / #4 / #5 / #6 / #7 / #8。")

    add_heading("1. 每节「四段式 + 图表点缀」结构", level=2)
    add_placeholder("（四段式模板说明）", height_lines=3)
    add_placeholder("（四段式模板说明）", height_lines=3)

    add_heading("2. 关键术语首次出现时给一句话定义", level=2)
    add_placeholder("（MoSCoW / RRF / HNSW / A3 等术语定义）", height_lines=4)

    add_heading("3. 表格前后各加一段叙述", level=2)
    add_placeholder("（表前说明 + 表后结论 · 落实批注 #2）", height_lines=3)

    add_heading("4. 图表必须有「图 X-Y 标题」", level=2)
    add_placeholder("（图编号规则 · 落实批注 #5）", height_lines=3)

    add_heading("5. 不使用 emoji（落实批注 #8）", level=2)
    add_placeholder("（禁用清单 + 替换为「是 / 否 / 满足 / 未满足」）", height_lines=3)

    add_heading("6. 1.5 文档组织结构（落实批注 #4）", level=2)
    add_placeholder("（3 段文字 + 1 张简图）", height_lines=3)

    add_page_break()

    # 目录（占位）
    add_heading("目录", level=0)
    add_meta_line("本目录由 Word 域代码自动生成，正文填充完成后按 F9 刷新。")
    p = doc.add_paragraph()
    set_paragraph_format(p, LINE_SPACING_BODY, 12, 6,
                         WD_ALIGN_PARAGRAPH.LEFT)
    r = p.add_run("（按 Ctrl + 点击此处 → 选择「更新域」→ 「更新整个目录」可生成完整目录）")
    set_cn_font(r, BODY_FONT_ZH, BODY_FONT_EN, SZ_SMALL,
                italic=True, color=GRAY_MID)

    add_page_break()

    # 图目录 + 表目录
    add_heading("图目录", level=1)
    add_meta_line("同目录 · 由 Word 域代码生成。")
    add_paragraph("", size=SZ_SMALL)

    add_heading("表目录", level=1)
    add_meta_line("同目录 · 由 Word 域代码生成。")
    add_paragraph("", size=SZ_SMALL)

    add_page_break()


# ============= 6. 正文：通用节构建函数 =============
def build_part_header(part_no_cn, part_name, target_pages, key_points):
    """部分标题页（part 起始分页）。"""
    add_heading(part_no_cn + "  " + part_name, level=0, page_break_before=True)
    add_target_box(target_pages, "中", "本部分为「" + part_name + "」")
    if key_points:
        add_meta_line("核心要点：")
        for kp in key_points:
            add_meta_line("  · " + kp)
    add_paragraph("", size=SZ_SMALL)


def build_h1(title, target_pages=1, priority="中", guidance=""):
    """一级标题（"1.1 …" 形式）。"""
    add_heading(title, level=1)
    add_target_box(str(target_pages) + " 页", priority, guidance)
    add_paragraph("", size=SZ_SMALL)


def build_h2(title, target_pages=None, priority="中", guidance=""):
    """二级标题。"""
    add_heading(title, level=2)
    if target_pages is not None:
        add_target_box(str(target_pages) + " 页", priority, guidance)


def build_h3(title, target_pages=None, priority="中", guidance=""):
    """三级标题。"""
    add_heading(title, level=3)
    if target_pages is not None:
        add_target_box(str(target_pages) + " 页", priority, guidance)


def build_four_para_placeholder(sec_id, sec_title, figures=None, tables=None):
    """四段式占位（铺垫/设计/取舍/表）· 落实批注 #6 / #7。"""
    labels = [
        ("段一 · 场景 / 痛点", "（1 段 · 约 200 字）", 4),
        ("段二 · 设计思路", "（1 段 · 约 200 字）", 4),
        ("段三 · 技术方案 / 关键决策", "（1-2 段 · 约 300-400 字）", 6),
        ("段四 · 价值与效果", "（1 段 · 约 200 字 · 量化或定性）", 4),
    ]
    for label, hint, height in labels:
        add_placeholder(sec_id + " " + label, hint, height_lines=height)
    # 表格
    if tables:
        for t in tables:
            add_table_placeholder(t["id"], t["header"], t["rows"], t["caption"])
    # 图
    if figures:
        for f in figures:
            add_figure_placeholder(f["id"], f["caption"], f.get("height", 6.0))


# ============= 7. 第一部分：作品概述 =============
def build_part1():
    build_part_header(
        "第一部分", "作品概述", 7,
        [
            "6-8 页 · 6 大版块的「是什么」",
            "项目背景 / 赛题理解 / 竞品 / 工作概要 / 文档结构",
            "覆盖 1.1-1.5 共 5 节",
        ]
    )

    # 1.1 项目背景与意义
    build_h1("1.1 项目背景与意义", 2, "高", "是什么 + 为什么")
    add_placeholder("1.1.1 高等教育个性化学习的现实困境",
                    hint="1 段铺垫 · 约 200-300 字", height_lines=5)
    add_placeholder("1.1.2 大模型与多智能体技术带来的新机遇",
                    hint="1 段说明技术拐点", height_lines=4)
    add_placeholder("1.1.3 项目研究与应用价值",
                    hint="1 段说明社会 / 学术 / 商业价值", height_lines=4)

    # 1.2 赛题理解
    build_h1("1.2 赛题理解", 1, "高", "A3 赛题约束")
    add_placeholder("1.2.1 A3 赛题核心约束",
                    hint="1 段铺垫 + 5 项必做的内在逻辑", height_lines=4)
    add_placeholder("1.2.2 必做项与加分项",
                    hint="表前 1 段说明 · 表 · 表后结论", height_lines=2)
    add_table_placeholder(
        "1-1", ["必做项", "本系统对应模块", "交付状态"], 5,
        "A3 赛题必做项 vs 本系统模块对应"
    )
    add_placeholder("1.2.3 非功能与合规要求",
                    hint="【重点改造节 · 落实批注 #0】四段式展开", height_lines=2)
    build_four_para_placeholder("1.2.3", "非功能与合规")
    add_table_placeholder(
        "1-2", ["指标", "目标值", "测量方式", "现状"], 5,
        "非功能指标表"
    )

    # 1.3 国内外研究现状与竞品分析
    build_h1("1.3 国内外研究现状与竞品分析", 2, "高", "现状 + 对比")
    add_placeholder("1.3.1 智能教育系统研究综述",
                    hint="1 段铺垫 · 2023-2026 拐点", height_lines=4)
    add_placeholder("1.3.2 多智能体在教育场景的应用",
                    hint="2-3 个学术工作 · 与本系统差异", height_lines=4)
    add_placeholder("1.3.3 RAG 与知识图谱在教育中的落地",
                    hint="引出 RAG + KG 双引擎", height_lines=4)
    add_placeholder("1.3.4 主流竞品横向对比表",
                    hint="【重点节 · 落实批注 #2】", height_lines=2)
    add_placeholder("1.3.4 段一", "对比逻辑（4 类 × 6 维 × 1 张主表）", 4)
    add_table_placeholder(
        "1-3",
        ["维度", "本系统", "传统在线教育", "通用 AI 对话"], 6,
        "主流竞品横向对比（学而思 AI / Speak / Quizlet Q-Chat / 豆包教育 / 文心一言教育 / 本系统）"
    )
    add_placeholder("1.3.4 段二", "表后结论 · 差异化护城河", 3)

    # 1.4 主要工作概要
    build_h1("1.4 主要工作概要", 1, "高", "12 模块 + 3 创新 + 关键数字")
    add_placeholder("1.4.1 系统能力一览",
                    hint="12 大核心模块 · 全部交付", height_lines=4)
    add_placeholder("1.4.2 三大创新点速览",
                    hint="RAG 引擎 / Agent 路由总线 / 五层 KG", height_lines=4)
    add_placeholder("1.4.3 核心指标摘要",
                    hint="P@5 / 引用准确率 / TTFT 等", height_lines=3)
    add_placeholder("1.4.4 关键规模数字",
                    hint="20+ 表 / 12 Agent / 7 资源类型 / 4 Provider", height_lines=3)

    # 1.5 文档组织结构
    build_h1("1.5 文档组织结构", 1, "中", "【落实批注 #4 · 3 段 + 1 张简图】")
    add_placeholder("1.5 段一", "六板块叙事主线", 4)
    add_placeholder("1.5 段二", "4 类读者推荐阅读路径（评委 / 教师 / 开发者 / 学生）", 5)
    add_placeholder("1.5 段三", "与答辩 PPT 17 张的双向对应", 4)
    add_figure_placeholder("图 1-1", "文档结构与 PPT 对应关系（三层映射）", 8.0)


# ============= 8. 第二部分：作品需求分析 =============
def build_part2():
    build_part_header(
        "第二部分", "作品需求分析", 8,
        [
            "7-9 页 · 6 大版块的「为什么」",
            "业务背景 / 用户旅程 / 功能性需求 / 非功能 / 需求映射",
            "覆盖 2.1-2.5 共 5 节",
        ]
    )

    # 2.1 业务背景与目标用户
    build_h1("2.1 业务背景与目标用户", 1, "中")
    add_placeholder("2.1.1 业务场景描述",
                    hint="3 段：学情差异 → 供给侧困境 → 系统定位", height_lines=6)
    add_figure_placeholder("图 2-1", "核心业务流程（注册→画像→生成→评估→回写）", 8.0)
    add_placeholder("2.1.2 用户角色定义",
                    hint="3 角色各 1 段 · 学生 / 教师 / 管理员", height_lines=6)

    # 2.2 用户旅程与场景画像
    build_h1("2.2 用户旅程与场景画像", 1, "中")
    add_placeholder("2.2.1 学生用户旅程图",
                    hint="典型一天 / 一周", height_lines=5)
    add_figure_placeholder("图 2-2", "学生用户旅程图", 8.0)
    add_placeholder("2.2.2 典型使用场景",
                    hint="预习 / 复习 / 应试 / 拓展 各 1 段", height_lines=6)

    # 2.3 功能性需求
    build_h1("2.3 功能性需求", 3, "高", "9 项功能 · 重点 2.3.2 / 2.3.8 / 2.3.9")
    add_placeholder("2.3.1 对话式学习画像构建（8 维）",
                    hint="8 维命名 + 更新策略", height_lines=4)
    add_placeholder("2.3.2 多智能体协同资源生成（7 种 + 3 入口）",
                    hint="【重点改造节 · 落实批注 #6 · 四段式】", height_lines=2)
    build_four_para_placeholder("2.3.2", "7 种资源 + 3 入口")
    add_table_placeholder(
        "2-1", ["#", "类型", "Agent", "关键特性"], 7,
        "7 种资源类型对照表"
    )
    add_placeholder("2.3.3 个性化学习路径规划", "KG 拓扑 + 薄弱点驱动", 3)
    add_placeholder("2.3.4 学习计划表生成", "collector / sequencer / scheduler / resource_linker", 3)
    add_placeholder("2.3.5 智能辅导", "RAG 问答 + B 站视频", 3)
    add_placeholder("2.3.6 学习效果评估", "答题闭环 + 画像回写", 3)
    add_placeholder("2.3.7 知识图谱可视化", "5 层节点 + 3 类关系", 3)
    add_placeholder("2.3.8 教学动画实时生成", "p5.js sketch + 沙箱", 3)
    add_placeholder("2.3.9 学习闭环", "4 阶段持续自适应", 3)

    # 2.4 非功能性需求
    build_h1("2.4 非功能性需求", 1, "中")
    add_placeholder("2.4.1 性能需求", "首字时延 / 并发 / 吞吐", 3)
    add_placeholder("2.4.2 可用性需求", "SLA / 错误恢复", 3)
    add_placeholder("2.4.3 安全与合规", "数据脱敏 / 内容审核", 3)
    add_placeholder("2.4.4 可维护性", "代码组织 / 文档", 3)
    add_placeholder("2.4.5 可扩展性", "新 Agent 接入 / 新 Provider", 3)

    # 2.5 需求优先级与评分项映射
    build_h1("2.5 需求优先级与评分项映射", 2, "高", "【2.5.1 重点改造节 · 落实批注 #7】")
    add_placeholder("2.5.1 MoSCoW 优先级矩阵",
                    hint="【四段式 + 表格前后叙述 + 术语定义】", height_lines=2)
    add_placeholder("2.5.1 段一", "MoSCoW 是什么 + 为什么选 MoSCoW", 4)
    add_placeholder("2.5.1 段二", "9 项需求如何映射到 4 档", 4)
    add_placeholder("2.5.1 段三", "MoSCoW 评估的边界", 3)
    add_placeholder("2.5.1 段四", "基于 MoSCoW 的版本规划 M1/M2/M3", 4)
    add_table_placeholder(
        "2-2", ["需求", "Must/Should/Could/Won't", "理由"], 9,
        "MoSCoW 优先级矩阵"
    )
    add_placeholder("2.5.1 表后结论", "X% 需求已 Must 级交付", 2)
    add_figure_placeholder("图 2-3", "需求优先级 vs 交付进度甘特图", 6.0)
    add_placeholder("2.5.2 需求 ↔ 赛题评分项对应表",
                    hint="表前 1 段 + 表 + 表后 1 段", height_lines=2)
    add_table_placeholder(
        "2-3", ["需求", "赛题评分项", "本系统实现"], 7,
        "需求 ↔ 赛题评分项对应表"
    )
    add_placeholder("2.5.3 需求 ↔ 实现状态追踪表", "已交付 / 部分 / 未交付", 2)
    add_table_placeholder(
        "2-4", ["需求", "实现状态", "负责人", "计划交付"], 9,
        "需求 ↔ 实现状态追踪表"
    )


# ============= 9. 第三部分：作品设计与实现（主体） =============
def build_part3():
    build_part_header(
        "第三部分", "作品设计与实现", 71,
        [
            "55-71 页 · 6 大版块的「怎么做」（文档主体）",
            "四层架构 / 12 Agent / 5 层防幻觉 / 7 资源 / 20+ 表 / Aurora UI",
            "覆盖 3.1-3.8 共 8 章",
        ]
    )

    # 3.1 总体设计
    build_h1("3.1 总体设计", 5, "高", "四层架构总览 + 技术栈")
    build_h2("3.1.1 设计目标与原则", 1)
    add_placeholder("3.1.1", "目标 + 原则", 4)

    build_h2("3.1.2 四层解耦总体架构", 1, "高")
    add_placeholder("3.1.2 段一", "四层架构的由来", 3)
    add_placeholder("3.1.2 段二", "层间单向依赖的设计动机", 3)
    add_placeholder("3.1.2 段三", "每层只解决一个问题", 3)
    add_figure_placeholder("图 3-1", "四层解耦总体架构（展示层 / 路由总线 / Agent 功能层 / 基础能力层）", 9.0)

    build_h2("3.1.3 技术栈选型", 1, "高")
    add_placeholder("3.1.3 段一", "技术选型三主线：异步 / 向量化 / 可切换", 4)
    add_table_placeholder(
        "3-1", ["类别", "技术选型", "选型理由"], 8,
        "技术栈一览（20+ 表 / 12 Agent / 7 资源 / 4 Provider）"
    )

    build_h2("3.1.4 运行环境", 1)
    add_placeholder("3.1.4 段一", "硬件最低配置", 2)
    add_placeholder("3.1.4 段二", "软件依赖与版本", 2)
    add_placeholder("3.1.4 段三", "启动方式", 2)

    build_h2("3.1.5 系统功能模块划分图", 1)
    add_figure_placeholder("图 3-2", "12 大核心模块全景图", 9.0)

    # 3.2 基础能力层
    build_h1("3.2 基础能力层：垂直领域模型基座引擎", 13, "高", "9 子节 · 创新点四（5.2）的技术底座")

    for sub_id, sub_title, sub_pages, hint in [
        ("3.2.1", "引擎定位与边界", 1, "业务痛点 / 设计目标 / 引擎边界"),
        ("3.2.2", "配置管理子系统", 1.5, "YAML + 环境变量 + 多 Provider + 提示词外置"),
        ("3.2.3", "LLM 接入子系统", 2, "Qwen3.6-Plus / Spark / 4 Provider / 重试 / 流式"),
        ("3.2.4", "Embedding 子系统", 1, "text-embedding-v4 1024 维"),
        ("3.2.5", "文档加载与解析子系统", 1.5, "PDF/DOCX/MD/TXT · 父子分块"),
        ("3.2.6", "向量索引子系统", 1.5, "pgvector + HNSW + tsvector"),
        ("3.2.7", "检索子系统", 2, "混合 / RRF / Query 改写 / 多样性"),
        ("3.2.8", "工具与外部能力集成", 1, "B 站视频检索 / 教学动画沙箱"),
        ("3.2.9", "基座引擎对外接口契约", 1, "LLM / Embedding / RAG / KG 接口"),
    ]:
        build_h2(sub_id + " " + sub_title, sub_pages)
        add_placeholder(sub_id, hint, height_lines=5)
        if sub_id in ["3.2.3", "3.2.7"]:
            add_table_placeholder(
                sub_id.replace(".", "-"),
                ["项目", "配置", "说明"], 4,
                sub_title + " 关键参数"
            )

    # 3.3 智能体中枢层
    build_h1("3.3 智能体中枢层：基于 LangGraph 的多智能体调度", 16, "高", "12 Agent + 3 模块 + 5 层防幻觉")

    build_h2("3.3.1 中枢层设计动机与边界", 1)
    add_placeholder("3.3.1", "动机 + 边界 + 3 模块组织", 5)

    build_h2("3.3.2 LangGraph 状态机设计", 1)
    add_placeholder("3.3.2", "StateGraph / AgentState / 12 节点 / 4 设计亮点", 5)
    add_figure_placeholder("图 3-3", "12 节点 LangGraph 状态机拓扑图", 9.0)

    # 3.3.3 十二智能体（每 Agent 1 页）
    build_h2("3.3.3 十二智能体详细设计", 12, "高", "每 Agent 独立 1 页")
    agents = [
        ("3.3.3.1", "profile_agent", "画像增量提取"),
        ("3.3.3.2", "planner_agent", "意图识别与路由"),
        ("3.3.3.3", "clarify_agent", "澄清问询"),
        ("3.3.3.4", "doc_agent", "文档生成"),
        ("3.3.3.5", "mindmap_agent", "思维导图"),
        ("3.3.3.6", "quiz_agent", "测验生成"),
        ("3.3.3.7", "code_agent", "代码示例"),
        ("3.3.3.8", "summary_agent", "总结生成"),
        ("3.3.3.9", "anim_agent", "教学动画生成（差异化亮点 · 1.5 页）"),
        ("3.3.3.10", "kg_agent", "知识图谱节点"),
        ("3.3.3.11", "safety_agent", "内容安全"),
        ("3.3.3.12", "recommend_agent", "资源推荐"),
    ]
    for ag_id, ag_name, ag_title in agents:
        pgs = 1.5 if "anim_agent" in ag_name else 1
        add_heading(ag_id + " " + ag_name + " — " + ag_title, level=3)
        add_target_box(str(pgs) + " 页", "中", "职责 / 输入输出 / 路由出口 / 提示词节选 / 示例")
        add_placeholder(ag_id + " 职责", "（1 段）", 2)
        add_placeholder(ag_id + " 输入输出", "（表格）", 2)
        add_table_placeholder(
            ag_id.replace(".", "-"),
            ["输入", "类型", "来源"], 4, ag_name + " 输入契约"
        )
        add_placeholder(ag_id + " 路由出口", "（1 段）", 2)
        add_placeholder(ag_id + " 提示词节选", "（代码块 · 系统提示词关键部分）", 3)
        add_placeholder(ag_id + " 示例", "（一段输入输出示例）", 3)

    build_h2("3.3.4 提示词工程体系", 1)
    add_placeholder("3.3.4", "NEVER/IMPORTANT/Do NOT/Avoid + Role/Rules/Pre-gen Check/Output", 5)

    # 3.3.5 五层防幻觉（4 页 · 重点）
    build_h2("3.3.5 五层递进防幻觉架构", 4, "高", "【重点专题】第一~五层 · 与 §4.6 消融实验呼应")
    add_placeholder("3.3.5 总述", "5 层递进 + 答 PPT 命名一致性", 4)
    add_figure_placeholder("图 3-4", "五层递进防幻觉总架构图", 9.0)

    for layer_id, layer_name, layer_title, hint in [
        ("3.3.5.1", "总问题分析", "防幻觉问题分析", "三类风险 / LLM 根本矛盾 / 教育底线"),
        ("3.3.5.2", "第一层", "检索前（EX-ANTE）— Query Rewrite", "指代消解 / 画像感知 / 多查询扩展"),
        ("3.3.5.3", "第二层", "检索中（RETRIEVAL）— 混合检索 + 父块回填", "混合 / RRF / 阈值 / 多样性"),
        ("3.3.5.4", "第三层", "生成中（IN-PROCESS）— NEVER/IMPORTANT", "分级禁令 / Pre-gen Check"),
        ("3.3.5.5", "第四层", "生成后（POST-HOC）— SafetyAgent", "4 项审核 / 只审不修 / fail-open"),
        ("3.3.5.6", "第五层", "全流程（CODE-LEVEL）— 硬约束 + LLM-as-Judge", "DB 校验 / 黄金回归 / 健康检查"),
        ("3.3.5.7", "效果度量", "防幻觉效果度量", "消融实验 / 关键指标"),
    ]:
        add_heading(layer_id + " " + layer_title, level=3)
        add_target_box("0.5 页", "高", hint)
        add_placeholder(layer_id, hint, height_lines=5)
        if "混合检索" in layer_title or "硬约束" in layer_title:
            add_table_placeholder(
                layer_id.replace(".", "-"),
                ["项", "内容"], 4, layer_title.split("—")[1].strip() + " 关键项"
            )

    build_h2("3.3.6 流式执行与进度可见性", 0.5)
    add_placeholder("3.3.6", "LangGraph stream + SSE", 3)

    build_h2("3.3.7 中枢层扩展性", 0.5)
    add_placeholder("3.3.7", "新增 Agent 标准化接入流程", 3)

    # 3.4 功能层
    build_h1("3.4 功能层：智能体驱动的服务模块", 11, "高", "11 个功能服务 · 重点 2/5/8/9/10")

    build_h2("3.4.1 对话式画像服务（8 维）", 1)
    add_placeholder("3.4.1", "8 维画像 + 完整度判断 + 快照", 4)
    add_table_placeholder("3-4-1", ["维度", "说明", "更新策略"], 8, "8 维学习画像")

    build_h2("3.4.2 资源生成服务（7 种 · 3 入口）", 1.5, "高")
    add_placeholder("3.4.2 段一", "资源生成的设计动机", 3)
    add_placeholder("3.4.2 段二", "7 种资源 + 3 入口的实现路径", 3)
    add_placeholder("3.4.2 段三", "异步任务编排的关键决策", 3)
    add_table_placeholder("3-4-2", ["#", "类型", "Agent", "关键特性"], 7, "7 种资源类型")
    add_table_placeholder("3-4-3", ["入口", "场景", "并发模型"], 3, "3 种生成入口")

    build_h2("3.4.3 资源库服务", 0.5)
    add_placeholder("3.4.3", "列表 / 详情 / 筛选 / 统计", 3)

    build_h2("3.4.4 学习路径服务", 0.5)
    add_placeholder("3.4.4", "CRUD / 条目 / 可视化", 3)

    build_h2("3.4.5 知识图谱服务（5 层 + 3 关系）", 1.5, "高")
    add_placeholder("3.4.5 段一", "5 层节点结构", 3)
    add_placeholder("3.4.5 段二", "3 类关系边", 3)
    add_placeholder("3.4.5 段三", "5 步自动构建流水线", 3)
    add_figure_placeholder("图 3-5", "5 层知识图谱自动构建流水线", 8.0)

    build_h2("3.4.6 智能辅导服务", 0.5)
    add_placeholder("3.4.6", "RAG 问答 + B 站视频", 3)

    build_h2("3.4.7 学习效果评估服务", 0.5)
    add_placeholder("3.4.7", "答题 + 掌握度 + 错题本", 3)

    build_h2("3.4.8 教学动画服务（PPT 差异化亮点）", 1, "高")
    add_placeholder("3.4.8 段一", "LLM 即时生成 vs 预录视频", 3)
    add_placeholder("3.4.8 段二", "anim-runtime.js 沙箱架构", 3)
    add_placeholder("3.4.8 段三", "工具集 u 与示例场景", 3)
    add_figure_placeholder("图 3-6", "教学动画运行时架构", 7.0)

    build_h2("3.4.9 学习计划表服务（4 步流水线）", 1, "高")
    add_placeholder("3.4.9", "collector / sequencer / scheduler / resource_linker", 4)
    add_figure_placeholder("图 3-7", "4 步学习计划表流水线", 6.0)

    build_h2("3.4.10 学习闭环（4 阶段）", 1, "高")
    add_placeholder("3.4.10", "路径 → 计划表 → 记录 → 画像回写", 4)
    add_figure_placeholder("图 3-8", "4 阶段学习闭环", 7.0)

    build_h2("3.4.11 用户认证与账号服务", 0.5)
    add_placeholder("3.4.11", "JWT + 邮箱验证 + 密码重置", 3)

    # 3.5 数据与接口设计
    build_h1("3.5 数据与接口设计", 9, "高", "20+ 表 + RESTful + Pydantic v2")

    build_h2("3.5.1 设计原则", 0.5)
    add_placeholder("3.5.1", "ORM 单一可信源 + 关系型+向量型 + Alembic + RESTful", 4)

    build_h2("3.5.2 数据库设计（20+ 张表）", 4, "高")
    add_placeholder("3.5.2 段一", "ER 图设计动机", 3)
    add_figure_placeholder("图 3-9", "数据库 ER 图（整页）", 12.0)
    table_groups = [
        ("3-5-2-1", "用户与会话", ["表名", "字段", "关系"], 3),
        ("3-5-2-2", "画像", ["表名", "字段", "关系"], 2),
        ("3-5-2-3", "知识图谱", ["表名", "字段", "关系"], 5),
        ("3-5-2-4", "资源", ["表名", "字段", "关系"], 4),
        ("3-5-2-5", "学习路径", ["表名", "字段", "关系"], 3),
        ("3-5-2-6", "学习计划表", ["表名", "字段", "关系"], 3),
        ("3-5-2-7", "答题", ["表名", "字段", "关系"], 3),
        ("3-5-2-8", "文档与向量", ["表名", "字段", "关系"], 2),
        ("3-5-2-9", "工具与外部", ["表名", "字段", "关系"], 3),
        ("3-5-2-10", "评估与监控", ["表名", "字段", "关系"], 3),
    ]
    for t_id, t_name, header, rows in table_groups:
        add_heading("数据表组：" + t_name, level=3)
        add_table_placeholder(t_id, header, rows, t_name + " 字段详表")

    build_h2("3.5.3 API 接口设计", 2)
    add_placeholder("3.5.3 段一", "接口设计规范（RESTful / Pydantic v2 / JWT）", 3)
    add_table_placeholder(
        "3-5-3-1", ["接口路径", "方法", "功能", "鉴权"], 12,
        "API 接口分组清单"
    )
    add_placeholder("3.5.3 关键接口", "POST /chat · /generate · /quiz · 详解", 2)
    add_heading("关键接口示例", level=3)
    add_placeholder("3.5.3.1", "POST /chat/{session_id} — 流式 SSE", 3)
    add_placeholder("3.5.3.2", "POST /generate + GET /generate/{task_id}/status", 3)
    add_placeholder("3.5.3.3", "GET /resources/{id}/quiz + POST /quiz/submit", 3)

    build_h2("3.5.4 错误处理与统一日志", 1.5)
    add_placeholder("3.5.4 段一", "错误码全表", 3)
    add_table_placeholder("3-5-4-1", ["错误码", "含义", "处理建议"], 8, "错误码全表")
    add_placeholder("3.5.4 段二", "loguru 分级 + trace_id 全链路", 3)

    # 3.6 前端
    build_h1("3.6 前端设计与实现（Aurora UI · 12 页面）", 5, "中")
    build_h2("3.6.1 设计目标与页面清单（12 个页面）", 1.5)
    add_placeholder("3.6.1 段一", "Aurora UI 设计语言", 3)
    add_placeholder("3.6.1 段二", "12 页面清单", 3)
    add_table_placeholder("3-6-1", ["#", "页面", "核心功能"], 12, "Aurora UI 12 页面清单")

    build_h2("3.6.2 关键交互设计", 1.5)
    add_placeholder("3.6.2.1", "流式输出与打字机效果", 2)
    add_placeholder("3.6.2.2", "Markdown 渲染 / 代码高亮", 2)
    add_placeholder("3.6.2.3", "思维导图渲染", 2)
    add_placeholder("3.6.2.4", "知识图谱渲染（ECharts / Cytoscape）", 2)
    add_placeholder("3.6.2.5", "动画演示沙箱（p5.js + iframe）", 2)

    build_h2("3.6.3 前端工程化", 1)
    add_placeholder("3.6.3", "原生 HTML/CSS/JS 选型理由 + 公共组件", 4)

    build_h2("3.6.4 可访问性与体验", 1)
    add_placeholder("3.6.4", "键盘 / 屏幕阅读器 / 弱网 / 移动端", 3)

    # 3.7 工程实践与部署
    build_h1("3.7 工程实践与部署", 4, "中")
    build_h2("3.7.1 性能优化", 1)
    add_placeholder("3.7.1", "全链路异步 + 缓存 + SSE + 双档", 4)

    build_h2("3.7.2 部署架构与启动流程", 1)
    add_placeholder("3.7.2", "Docker pgvector + conda + uvicorn + Alembic", 4)
    add_figure_placeholder("图 3-10", "部署架构图", 7.0)

    build_h2("3.7.3 监控、日志与备份恢复", 1)
    add_placeholder("3.7.3", "日志切分 + trace_id + 备份策略 + L1 健康检查", 4)

    build_h2("3.7.4 项目管理与开发流程", 1)
    add_placeholder("3.7.4", "Git Flow + 提交规范 + 团队分工 + 答辩配套", 4)

    # 3.8 端到端案例
    build_h1("3.8 端到端典型案例：从画像构建到完整学习路径", 4, "高", "评委爱看的故事化演示")
    build_h2("3.8.1 场景设定（虚拟用户「张同学」）", 0.5)
    add_placeholder("3.8.1", "张同学 · 计算机专业 · 复习数据结构", 4)

    build_h2("3.8.2 全流程时序图（12 步串起 12 个 Agent）", 1)
    add_figure_placeholder("图 3-11", "张同学端到端时序图", 12.0)

    build_h2("3.8.3 关键交互截图（8-10 张）", 1.5)
    for i in range(1, 6):
        add_figure_placeholder("图 3-12-" + str(i),
                               "关键交互截图 #" + str(i), 6.0)

    build_h2("3.8.4 各模块在该案例中的协作分工", 1)
    add_placeholder("3.8.4", "12 Agent + 7 资源 + 4 阶段闭环的协作", 4)


# ============= 10. 第四部分：作品测试与分析 =============
def build_part4():
    build_part_header(
        "第四部分", "作品测试与分析", 10,
        [
            "8-10 页 · 6 大版块的「做得怎么样」",
            "测试体系 / RAG 四层评估 / 黄金集 / 五层消融",
            "覆盖 4.1-4.8 共 8 节",
        ]
    )

    for sec_id, sec_title, sec_pages, hint in [
        ("4.1", "测试体系总览", 0.5, "单元 / 集成 / 评估三层"),
        ("4.2", "单元与集成测试", 1, "pytest-asyncio + 覆盖率"),
        ("4.3", "RAG 四层评估架构", 1, "L1 健康 / L2 检索 / L3 生成 / L4 黄金"),
        ("4.4", "黄金数据集设计与实验", 1.5, "15 条人工标注 + expected_aspects"),
        ("4.5", "评估指标定义与结果", 1.5, "公式 + 优化前后对比"),
        ("4.6", "五层防幻觉效果分析", 2, "整体指标 + 消融实验 + Bad Case"),
        ("4.7", "性能测试", 1, "TTFT / 端到端 / 并发 / 资源"),
        ("4.8", "优化迭代闭环", 1, "评估 → 诊断 → 改进 → 复评"),
    ]:
        build_h1(sec_id + " " + sec_title, sec_pages)
        add_placeholder(sec_id, hint, height_lines=5)

    # 4.3 表格
    add_heading("4.3 RAG 四层评估表", level=2)
    add_table_placeholder("4-3-1", ["层", "关注", "关键指标", "成本"],
                          4, "RAG 四层评估架构（L1-L4）")

    # 4.6 消融实验
    add_heading("4.6.1 消融实验", level=2)
    add_table_placeholder("4-6-1",
                          ["配置", "幻觉率", "引用准确率", "误导率"],
                          6, "五层防幻觉消融实验（每层关闭后）")
    add_figure_placeholder("图 4-1", "各层贡献度对比柱状图", 7.0)

    # 4.7 性能
    add_heading("4.7 性能测试表", level=2)
    add_meta_line("（落实批注 #8 · 表头与正文不使用 emoji · 用「是 / 否 / 满足 / 未满足」）")
    add_table_placeholder("4-7-1",
                          ["是否满足", "指标", "目标值", "实测值", "备注"],
                          6, "性能指标达成情况")

    # 4.8 优化闭环
    add_heading("4.8 优化案例", level=2)
    for i in range(1, 4):
        add_placeholder("4.8." + str(i), "优化案例 #" + str(i), 4)


# ============= 11. 第五部分：作品创新性说明 =============
def build_part5():
    build_part_header(
        "第五部分", "作品创新性说明", 7,
        [
            "6-8 页 · 6 大版块的「凭什么领先」",
            "3 PPT 创新点 + 1 文档独有创新点 + 综合分析",
            "覆盖 5.1-5.5",
        ]
    )

    build_h1("5.1 创新点一（PPT）：教材驱动的 RAG 知识库引擎（5 层防幻觉闸门）", 1.5, "高")
    add_placeholder("5.1.1", "问题陈述：通用 LLM 垂直课程幻觉严重", 3)
    add_placeholder("5.1.2", "技术方案：教材权威 + 5 层防幻觉", 3)
    add_placeholder("5.1.3", "检索流水线 4 步", 3)
    add_placeholder("5.1.4", "五层防幻觉闸门（第一~五层）", 3)
    add_placeholder("5.1.5", "四层评估体系（L1-L4）", 3)
    add_placeholder("5.1.6", "实测效果：P@5 / 引用准确率 90%+", 3)
    add_figure_placeholder("图 5-1", "RAG 引擎 + 5 层防幻觉总架构", 7.0)

    build_h1("5.2 创新点四（文档独有 · PPT 不讲）：垂直领域模型基座引擎", 1.5, "中",
             "【落实批注 #9】答辩时仅在被追问时提及")
    add_placeholder("5.2.1", "问题陈述：多模型多工具散落业务代码", 3)
    add_placeholder("5.2.2", "设计思路：LLM/Embedding/RAG/工具 打包为基座", 3)
    add_placeholder("5.2.3", "技术方案 5 子系统", 3)
    add_placeholder("5.2.4", "核心亮点：配置外置 / 提示词外置 / 异常降级 / 解耦", 3)
    add_placeholder("5.2.5", "价值量化：新增 Provider 1 天 / 新增 Agent 0.5 天", 3)
    add_placeholder("5.2.6", "与 3 个 PPT 创新点的关系（使能层）", 3)
    add_placeholder("5.2.7", "与 §3.2 基础能力层的对应", 2)

    build_h1("5.3 创新点二（PPT）：智能体路由总线与模块化协同（12 Agent · 3 模块 · 4 亮点）", 1.5, "高")
    add_placeholder("5.3.1", "问题陈述：多智能体分散调用、状态难管理", 3)
    add_placeholder("5.3.2", "技术方案：LangGraph 状态机 + 12 节点", 3)
    add_placeholder("5.3.3", "4 项设计亮点（共享状态 / 条件路由 / DB 注入 / 画像驱动）", 4)
    add_placeholder("5.3.4", "3 模块组织（入口 / 生成 / 出口）", 3)
    add_placeholder("5.3.5", "价值", 3)
    add_figure_placeholder("图 5-2", "智能体路由总线 · 4 亮点示意", 7.0)

    build_h1("5.4 创新点三（PPT）：五层知识图谱自动构建（零人工标注 · 增量更新）", 1.5, "高")
    add_placeholder("5.4.1", "问题陈述：课程知识体系人工维护成本高", 3)
    add_placeholder("5.4.2", "技术方案：PDF → 结构化知识网络", 3)
    add_placeholder("5.4.3", "5 层节点结构", 3)
    add_placeholder("5.4.4", "3 类关系边", 3)
    add_placeholder("5.4.5", "5 步构建流水线", 3)
    add_placeholder("5.4.6", "应用场景（路径 / 推荐 / 计划表）", 3)
    add_placeholder("5.4.7", "价值", 2)
    add_figure_placeholder("图 5-3", "5 层 KG 自动构建流水线", 7.0)

    build_h1("5.5 创新点综合价值与差异化分析", 1, "高")
    add_placeholder("5.5 段一", "3 PPT + 1 文档独有 的协同效应", 3)
    add_table_placeholder("5-5-1",
                          ["维度", "本系统", "传统在线教育", "通用 AI 对话"],
                          6, "6 维度竞品对比")
    add_placeholder("5.5 段二", "对教育 AI 落地的护城河", 3)
    add_placeholder("5.5 段三", "文档独有第 4 创新点的展示价值", 2)


# ============= 12. 第六部分：作品总结 =============
def build_part6():
    build_part_header(
        "第六部分", "作品总结", 4,
        [
            "3-5 页 · 6 大版块的「成果与未来」",
            "12 模块完成度 + 应用价值 + 4 方向 Roadmap",
            "覆盖 6.1-6.3",
        ]
    )

    build_h1("6.1 工作总结", 1.5, "高", "12 个核心模块 · 100% 完成")
    add_placeholder("6.1.1", "12 模块完成度（对照 §2.5 需求矩阵）", 3)
    add_table_placeholder("6-1-1", ["#", "功能模块", "状态", "备注"], 12, "12 模块完成度")
    add_placeholder("6.1.2", "已完成内容核对", 3)
    add_placeholder("6.1.3", "关键技术成果归纳", 3)

    build_h1("6.2 应用价值", 1, "中")
    add_placeholder("6.2.1", "教育场景社会价值", 3)
    add_placeholder("6.2.2", "工程方法论价值", 3)
    add_placeholder("6.2.3", "商业化潜力", 3)

    build_h1("6.3 不足与未来工作", 1.5, "中", "4 方向 Roadmap（与 PPT 一致）")
    add_placeholder("6.3.1", "当前局限", 3)
    add_placeholder("6.3.2", "短期优化方向", 3)
    add_placeholder("6.3.3 方向 1", "多课程知识库", 3)
    add_placeholder("6.3.3 方向 2", "移动端适配（小程序 / PWA）", 3)
    add_placeholder("6.3.3 方向 3", "协作学习（互评 / 互测 / 互讲）", 3)
    add_placeholder("6.3.3 方向 4", "更精细的自适应（tracker.js + 注意力建模）", 3)


# ============= 13. 附录 =============
def build_appendix():
    add_heading("附录", level=0, page_break_before=True)
    add_target_box("8-12 页", "低", "附录部分 · 不计入正文 80-100 页目标")

    for app_id, app_title, hint in [
        ("A", "API 接口完整清单", "全部接口 · 请求/响应示例"),
        ("B", "数据库表结构详细字段", "20+ 表 · 字段 / 类型 / 索引 / 关系"),
        ("C", "系统截图集", "12 页面 + 关键交互 · 50+ 张截图"),
        ("D", "AI 辅助工具使用说明", "讯飞星辰 / AI 写作 / 提示词审计"),
    ]:
        build_h1("附录 " + app_id + " " + app_title, 2, "低")
        add_placeholder("附录" + app_id, hint, height_lines=5)

    add_heading("参考文献", level=1, page_break_before=True)
    add_placeholder("参考文献", "（约 20-30 条 · GB/T 7714 格式）", 10)

    add_heading("致谢", level=1)
    add_placeholder("致谢", "（约 1 段 · 致谢指导教师 / 团队 / 评委）", 4)


# ============= 14. 末节：篇幅分配总表 + 差异说明 =============
def build_tail():
    add_heading("篇幅分配总表（v3.2 模板目标）", level=0, page_break_before=True)
    add_meta_line("下表为正式撰写时按本模板填充的篇幅分配，正文目标 80-100 页。")
    add_table_placeholder("末-1", ["板块", "章节", "预计页数", "状态"], 20,
                          "篇幅分配总表（按 v3.2 重排）")

    add_heading("与 v2 / v3.0 / v3.1 / v3.2 的差异说明", level=0)
    add_meta_line("本节是模板与历次版本的差异自检，撰写时同步更新。")
    for sec in [
        "v3.1 → v3.2 · 落实 YC 2026.06.27 10 条批注 · 评审批注闭环",
        "v3.0 → v3.1 · 与答辩 PPT 对齐（项目名 / 4 层架构 / 12 Agent / 5 层防幻觉 / 4 步 KG / 4 阶段闭环 / 4 Provider / 4 Roadmap）",
        "v2 → v3 整合 · 11 章合并为 6 大版块 · 创新点从附录提到正文",
    ]:
        add_placeholder(sec, "（差异表 · 详见 docs/xqt/文档设计/设计文档大纲_v3.md 末节）", 6)

    add_heading("模板使用说明（给后续撰写人）", level=1)
    add_meta_line("1) 撰写人应先按 [篇幅目标] 估算每节字数；")
    add_meta_line("2) 落实四段式（场景/痛点 → 设计 → 取舍 → 价值）；")
    add_meta_line("3) 表格前后加 1 段叙述；")
    add_meta_line("4) 图统一编号「图 X-Y 标题」并加图说；")
    add_meta_line("5) 全文不出现 emoji，统一用「是 / 否 / 满足 / 未满足」；")
    add_meta_line("6) 关键术语首次出现给一句话定义。")

    add_meta_line("")
    add_meta_line("模板结束。正文待撰写。")


# ============= 15. 主流程 =============
def main():
    build_cover()
    build_front_matter()
    build_part1()
    build_part2()
    build_part3()
    build_part4()
    build_part5()
    build_part6()
    build_appendix()
    build_tail()
    doc.save(OUTPUT_PATH)
    print("Saved:", OUTPUT_PATH)


if __name__ == "__main__":
    main()
